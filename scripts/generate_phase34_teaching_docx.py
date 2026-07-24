from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "phase3_phase4_teaching_guide.docx"


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    for name, size in [
        ("Title", 22),
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name].font.size = Pt(size)

    if "Code" not in styles:
        styles.add_style("Code", 1)
    styles["Code"].font.name = "Consolas"
    styles["Code"].font.size = Pt(9)
    styles["Code"].font.color.rgb = RGBColor(32, 32, 32)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str = "") -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def num(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Code")
    for i, line in enumerate(text.splitlines()):
        if i:
            para.add_run().add_break()
        para.add_run(line)


def table(doc: Document, headers, rows) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def add_title(doc: Document) -> None:
    doc.add_heading("Phase 3 + Phase 4 教学文档", 0)
    p(doc, "主题：计算图引擎、内存池、CPU 推理加速后端、图级并行、INT8 权重量化")
    p(doc, "目标：读完后能够知道每个文件负责什么、每个类和函数做什么、数据如何在 Phase3 和 Phase4 之间流动，以及为什么这些实现方式适合推理框架。")
    p(doc, "适用代码范围：phase3_graph/ 与 phase4_gpu/ 当前工作树。")
    p(doc, "阅读方法：先看总览，再看逐文件逐函数讲解，最后用执行链路把所有函数串起来。")


def add_overview(doc: Document) -> None:
    h(doc, "1. 总览：Phase3 和 Phase4 分别解决什么问题", 1)
    p(doc, "Phase3 解决的是“怎么表示和执行一个模型”。它把模型拆成很多 Node，每个 Node 表示一个算子；Tensor 存储数据；ComputeGraph 保存 DAG；JsonParser 从配置文件构建图；Executor 按拓扑顺序调度算子；MemoryPool 尝试复用中间 Tensor 的内存。")
    p(doc, "Phase4 解决的是“在没有 CUDA GPU 的情况下，怎么让推理更快”。它不推翻 Phase3，而是在 Phase3 的图结构上增加 CPU 加速层：Backend 抽象、OpenMP 算子后端、按 DAG level 并行执行、INT8 weight-only 量化 GEMM。")
    table(
        doc,
        ["阶段", "核心问题", "关键文件", "产物"],
        [
            ["Phase3", "模型如何被表示、排序、调度和管理内存", "graph/json_parser/executor/memory_pool", "可以执行 JSON 描述的计算图"],
            ["Phase4", "没有 CUDA 时如何做 CPU 推理加速", "backend/parallel_graph_executor/quantization", "CPU 并行后端、图级并行、INT8 量化路径"],
        ],
    )
    h(doc, "1.1 一条完整推理链路", 2)
    num(doc, "JSON 文件描述节点：q_proj、k_proj、v_proj、attn_out、ln_out。")
    num(doc, "JsonParser 把 JSON 变成 ComputeGraph 里的 Node。")
    num(doc, "用户或测试代码把输入 x 和权重 Wq/Wk/Wv 放入 ComputeGraph 的 tensor store。")
    num(doc, "ComputeGraph::topological_sort 得到合法执行顺序。")
    num(doc, "Phase3 Executor 或 Phase4 ParallelGraphExecutor 逐节点执行。")
    num(doc, "每个节点根据 op_type 调用 MatMul、Attention、LayerNorm 等算子。")
    num(doc, "输出 Tensor 写回 graph，成为后续节点输入。")
    code(doc, """JSON -> JsonParser -> ComputeGraph
ComputeGraph + input tensors -> topological_sort
sorted nodes -> Executor / ParallelGraphExecutor
node.op_type -> Backend or Phase2/Phase1 operator
output tensor -> graph.set_tensor -> next node""")


def add_phase3_graph(doc: Document) -> None:
    h(doc, "2. Phase3：graph.h / graph.cpp", 1)
    p(doc, "这组文件是计算图的核心数据结构。你可以把它理解为一个极简版 ONNX Runtime 图表示。")
    table(
        doc,
        ["对象/函数", "位置", "作用", "怎么和其他代码组合"],
        [
            ["Tensor", "graph.h", "保存一个命名张量：name、shape、data", "Executor 和 Backend 都围绕 Tensor 读写数据"],
            ["Tensor::total_size", "graph.h", "把 shape 每一维相乘，得到元素总数", "用于判断张量理论大小"],
            ["Tensor::zeros", "graph.h", "创建指定 shape 的 Tensor，并把 data 初始化为 0", "各算子创建输出 Tensor 时使用"],
            ["Node", "graph.h", "保存一个图节点：id、op_type、inputs、outputs、params", "拓扑排序和执行器都只看 Node，不关心具体模型"],
            ["ComputeGraph::add_node", "graph.cpp", "把节点放入 nodes_，并记录 id 到下标的索引", "JsonParser 和手写测试构图时调用"],
            ["ComputeGraph::add_weight", "graph.cpp", "把权重 Tensor 放入 tensor store", "功能上等价于 set_tensor，语义上表示权重"],
            ["ComputeGraph::find_node", "graph.cpp", "根据 id 找 Node", "测试或调试时验证图结构"],
            ["ComputeGraph::get_tensor", "graph.cpp", "根据名字找 Tensor，找不到返回 nullptr", "Executor 每次取输入都走这里"],
            ["ComputeGraph::set_tensor", "graph.cpp", "把 Tensor 写入 tensor store", "输入、权重、中间结果、最终输出都写这里"],
            ["ComputeGraph::topological_sort", "graph.cpp", "Kahn 算法产生合法执行顺序", "Executor 必须按这个顺序执行"],
            ["ComputeGraph::dump", "graph.cpp", "打印节点、输入输出、参数和 Tensor 数量", "调试图结构"],
        ],
    )
    h(doc, "2.1 Tensor：为什么只用 vector<float>", 2)
    p(doc, "当前项目是教学型推理框架，Tensor 只保留最少字段。name 用来在图里查找；shape 用来解释 data 的维度；data 是行主序的一维 float 数组。真实框架还会有 dtype、stride、device、layout、allocator 等信息，但 Phase3 暂时不需要。")
    code(doc, """struct Tensor {
    std::string name;
    std::vector<int> shape;
    std::vector<float> data;
};""")
    p(doc, "注意：data 是一维数组，不是二维数组。例如 shape=[128,64] 时，data[i*64+j] 表示第 i 行第 j 列。")
    h(doc, "2.2 Node：计算图中的一个算子调用", 2)
    p(doc, "Node 不保存真实数据，只保存“我要对哪些 Tensor 做什么操作，并把结果叫什么”。这正是计算图的核心：图结构和数据分离。")
    code(doc, """Node{
  id      = "q_proj",
  op_type = "MatMul",
  inputs  = {"x", "Wq"},
  outputs = {"q_proj"},
  params  = {}
}""")
    h(doc, "2.3 topological_sort：为什么执行前必须排序", 2)
    p(doc, "如果一个节点依赖另一个节点的输出，它必须后执行。例如 attn_out 依赖 q_proj、k_proj、v_proj。拓扑排序就是把 DAG 排成一个合法线性序列。")
    code(doc, """q_proj  k_proj  v_proj     level 0
       \\    |    /
        attn_out             level 1
            |
          ln_out             level 2""")
    p(doc, "算法核心是 in-degree。某节点的 in-degree 表示它还缺多少上游节点。in-degree 为 0 的节点可以先执行。每执行完一个节点，就把依赖它的节点 in-degree 减 1。")
    num(doc, "建立 name_to_node，当前代码用节点 id 代表它默认输出的名字。")
    num(doc, "给每个节点初始化 in_degree=0。")
    num(doc, "扫描每个节点的 inputs，如果 input 是某个节点产物，就给当前节点 in_degree++。")
    num(doc, "把 in_degree=0 的节点放进队列。")
    num(doc, "不断弹出队列头，加入 sorted；再更新消费者的 in_degree。")
    num(doc, "如果 sorted 数量小于节点数，说明图里有环，返回空。")


def add_phase3_json(doc: Document) -> None:
    h(doc, "3. Phase3：json_parser.h / json_parser.cpp", 1)
    p(doc, "JsonParser 是把文件形式的模型配置变成 ComputeGraph 的入口。它是手写解析器，只支持本项目需要的 JSON 子集。")
    table(
        doc,
        ["函数/类", "作用", "关键点"],
        [
            ["Tokenizer::peek", "跳过空白后看当前字符，但不移动位置", "用于 while(tok.peek()!=...) 判断结构是否结束"],
            ["Tokenizer::next", "跳过空白后取当前字符，并移动 pos_", "解析符号时使用"],
            ["Tokenizer::expect", "要求下一个字符必须是指定字符，否则抛异常", "保证 JSON 结构合法"],
            ["Tokenizer::read_string", "读取双引号包围的字符串", "解析 key、id、op、inputs"],
            ["Tokenizer::read_number", "读取数字并转 float", "解析 params 里的 eps 等参数"],
            ["Tokenizer::skip_whitespace", "跳过空格、换行、tab 和逗号", "简化解析循环"],
            ["parse_node", "解析一个节点对象并 graph.add_node", "如果 outputs 为空，默认 outputs[0]=id"],
            ["JsonParser::parse", "解析完整 JSON 字符串", "只关心 nodes 字段，异常时返回 false"],
            ["JsonParser::parse_file", "读文件内容后调用 parse", "测试里用它读取 configs/attention_block.json"],
        ],
    )
    h(doc, "3.1 parse_node 的工作方式", 2)
    p(doc, "parse_node 每次看到一个 { ... }，就构造一个 Node。它循环读取 key，根据 key 决定把后面的 value 放到 node 的哪个字段。")
    code(doc, """if (key == "id")        node.id = tok.read_string();
else if (key == "op")   node.op_type = tok.read_string();
else if (key == "inputs")  读取字符串数组;
else if (key == "params")  读取 key=float 的对象;""")
    p(doc, "项目 JSON 里没有显式 outputs，所以 parse_node 最后会把 outputs 自动设成 id。这就是为什么 q_proj 节点的输出 Tensor 名也叫 q_proj。")
    h(doc, "3.2 parse 的容错边界", 2)
    p(doc, "这个解析器故意很小，不支持布尔值、null、转义字符串、科学计数法等完整 JSON 特性。它适合本项目教学，但不是生产级 JSON parser。生产代码应使用 nlohmann/json、RapidJSON 或 simdjson。")


def add_phase3_memory(doc: Document) -> None:
    h(doc, "4. Phase3：memory_pool.h / memory_pool.cpp", 1)
    p(doc, "MemoryPool 解决的是推理过程中频繁分配中间 Tensor 的问题。它预先申请一大块连续内存，然后让不同 Tensor 借用其中一段。")
    table(
        doc,
        ["函数/成员", "作用", "细节"],
        [
            ["Slot", "记录池中一个切片", "name 用于调试，ptr 指向切片开头，size 是 float 数量，in_use 表示是否被占用"],
            ["MemoryPool::MemoryPool", "申请大 buffer", "Windows 用 _aligned_malloc，其他平台用 aligned_alloc；32 字节对齐为 AVX2 准备"],
            ["MemoryPool::~MemoryPool", "释放大 buffer", "Windows 用 _aligned_free，其他平台 free"],
            ["allocate", "分配 count 个 float", "先复用空闲 slot，再追加到尾部，不够时尝试 compaction"],
            ["deallocate", "按 name 标记 slot 空闲", "不归还 OS，只是允许后续复用"],
            ["base", "返回底层 pool_ 指针", "低层算子可直接访问"],
            ["used_size", "统计正在使用的 float 数", "遍历 slots_ 累加 in_use slot"],
            ["free_size", "total_size - used_size", "快速看剩余容量"],
            ["peak_used", "历史峰值使用量", "用于评估内存峰值"],
            ["allocation_count", "累计分配次数", "用于观察复用是否有效"],
            ["dump", "打印 pool 状态", "调试内存复用过程"],
        ],
    )
    h(doc, "4.1 allocate 的两阶段策略", 2)
    num(doc, "先找已经释放的 slot，如果大小足够，就直接复用。")
    num(doc, "如果没有合适 slot，就计算当前所有 slot 的最大尾部 offset，在末尾追加新 slot。")
    num(doc, "如果末尾空间不够，就把所有 in_use slot 向前 memmove，形成连续空闲尾部。")
    num(doc, "如果压缩后仍然不够，返回 nullptr，表示池容量不足。")
    code(doc, """free slot reuse:
[A in use][B free][C in use] -> D can reuse B if size fits

append:
[A in use][C in use][free...] -> new D at the end

compact:
[A in use][free][C in use][free...] -> [A][C][free...]""")
    h(doc, "4.2 当前 MemoryPool 的教学限制", 2)
    p(doc, "Phase3 的 Tensor 仍然持有 std::vector<float>，所以 pool 路径里经常会发生 pool buffer 和 vector 之间的拷贝。这对教学有价值：你能看到生命周期分析和复用逻辑；但它还不是零拷贝的生产级内存池。Phase4 文档里也把这个作为后续优化点。")


def add_phase3_executor(doc: Document) -> None:
    h(doc, "5. Phase3：executor.h / executor.cpp", 1)
    p(doc, "Executor 是 Phase3 的运行时。它拿到拓扑排序后的 Node 列表，逐个根据 op_type 调用对应算子。")
    table(
        doc,
        ["函数/结构", "作用", "执行时机"],
        [
            ["require_tensor", "从 graph 取输入 Tensor，缺失时打印错误", "每个算子执行前"],
            ["TensorLifetime", "记录一个 Tensor 的 first_use、last_use、is_input", "pool 执行前生命周期分析"],
            ["analyze_lifetimes", "扫描 sorted_nodes，计算每个 Tensor 最后一次被谁使用", "execute_with_pool 开始时"],
            ["execute_matmul", "执行 MatMul，调用 Phase1 gemm_naive_ikj", "普通执行路径"],
            ["execute_softmax", "复制输入 Tensor 后逐行 softmax_2d", "普通执行路径"],
            ["execute_layernorm", "对整个 flattened Tensor 做 layernorm", "普通执行路径"],
            ["execute_gelu", "调用 gelu_vector_exact", "普通执行路径"],
            ["execute_attention", "调用 Phase2 scaled_dot_product_attention", "普通执行路径"],
            ["execute_matmul_pool", "从 MemoryPool 分配输出空间后执行 MatMul", "pool 执行路径"],
            ["execute_softmax_pool", "pool 中分配输出，复制输入后原地 softmax", "pool 执行路径"],
            ["execute_layernorm_pool", "pool 中分配输出后 layernorm", "pool 执行路径"],
            ["execute_gelu_pool", "pool 中分配输出后 GELU", "pool 执行路径"],
            ["execute_attention_pool", "计算 attention 输出后复制进 pool", "pool 执行路径"],
            ["Executor::execute", "按 sorted_nodes 线性执行普通路径", "Phase3 默认执行"],
            ["Executor::execute_with_pool", "按 sorted_nodes 线性执行 pool 路径，并释放过期 Tensor", "Phase3 内存池执行"],
            ["Executor::get_output", "从 graph 获取输出 Tensor", "测试和调用者取结果"],
        ],
    )
    h(doc, "5.1 普通执行路径：以 MatMul 为例", 2)
    p(doc, "MatMul 节点的 inputs 通常是 {\"x\", \"Wq\"}。execute_matmul 先拿到 A/B，然后根据 shape 推出 M、K、N，检查 A 的列数是否等于 B 的行数。通过后创建输出 Tensor C，并调用 Phase1 的 GEMM。")
    code(doc, """A shape = [M, K]
B shape = [K, N]
C shape = [M, N]
C = A x B""")
    p(doc, "这个函数体现了图执行器的典型职责：检查输入、推导输出 shape、调用底层算子、把输出写回 graph。")
    h(doc, "5.2 Softmax / GELU / LayerNorm 的共性", 2)
    p(doc, "这些算子都是一进一出。它们先通过 require_tensor 取输入，再创建输出，然后调用 Phase2 算子库。Softmax 因为 Phase2 的 softmax_2d 是原地修改，所以先复制输入再处理。GELU 使用 exact 版本，强调正确性。")
    p(doc, "Phase3 的 LayerNorm 使用 N=in->data.size()，也就是把整个 Tensor 展平成一个向量来归一化。这适合早期验证，但 Transformer 推理通常应该按每个 token 的 hidden dimension 逐行归一化。Phase4 的 CpuParallelBackend 已改成 row-wise LayerNorm。")
    h(doc, "5.3 Attention 的组合", 2)
    p(doc, "execute_attention 不是自己写 QK^T/Softmax/WV，而是把 Q/K/V 传给 Phase2 的 scaled_dot_product_attention。这样 Phase3 不关心 Attention 细节，只负责调度。")
    code(doc, """Q = q_proj
K = k_proj
V = v_proj
scores  = Q x K^T / sqrt(d_k)
weights = softmax(scores)
out     = weights x V""")
    h(doc, "5.4 execute 的主循环", 2)
    p(doc, "Executor::execute 是最直观的执行方式：for 每个 node，根据 op_type 进入 if-else 分发。每个算子成功后，它会从 graph 取输出 Tensor 并打印 shape。")
    code(doc, """for node in sorted_nodes:
    if op_type == "MatMul": execute_matmul
    if op_type == "Softmax": execute_softmax
    if op_type == "LayerNorm": execute_layernorm
    if op_type == "GELU": execute_gelu
    if op_type == "Attention": execute_attention""")
    p(doc, "这种硬编码 if-else 是教学阶段最容易理解的做法。Phase4 的 Backend 抽象就是对它的升级：把“怎么执行某个 op”从 Executor 里抽出去。")
    h(doc, "5.5 execute_with_pool 的主循环", 2)
    p(doc, "execute_with_pool 在执行前调用 analyze_lifetimes。每执行完一个节点，就检查哪些 Tensor 的 last_use 等于当前节点下标，并调用 pool.deallocate。这样后续 Tensor 可以复用这段内存。")
    p(doc, "理解这个函数的关键：拓扑顺序不仅决定谁先算，还给生命周期分析提供了时间轴。idx 越小表示越早执行，last_use 就表示“这个 Tensor 最后被哪个节点读”。")


def add_phase4_backend(doc: Document) -> None:
    h(doc, "6. Phase4：backend.h / cpu_parallel_backend.cpp", 1)
    p(doc, "Phase4 的第一件事是把算子执行从 Phase3 Executor 的硬编码 if-else 中抽象出来。Backend 是一个 execution provider 接口，CpuParallelBackend 是当前没有 CUDA GPU 时的 CPU 加速实现。")
    table(
        doc,
        ["函数/类", "作用", "为什么重要"],
        [
            ["Backend::~Backend", "虚析构函数", "通过基类指针删除派生类时行为正确"],
            ["Backend::name", "返回后端名字", "日志、benchmark、调试可区分 backend"],
            ["Backend::execute", "执行一个 Node", "把 op_type 分发给具体后端实现"],
            ["CpuParallelBackend::name", "返回 cpu-parallel", "标识当前是 CPU 并行后端"],
            ["CpuParallelBackend::execute", "按 op_type 调用后端内部函数", "Phase4 的算子入口"],
            ["require_input_count", "检查输入数量", "提前发现图配置错误"],
            ["execute_matmul", "调用 gemm_parallel_tiled", "Phase4 最主要的 CPU 性能来源"],
            ["execute_softmax", "逐行并行 softmax", "Attention scores 每一行互不依赖"],
            ["execute_layernorm", "逐行并行 LayerNorm", "符合 Transformer 按 hidden 维归一化的语义"],
            ["execute_gelu", "并行 GELU approx", "推理场景用近似版提高速度"],
            ["execute_attention", "用并行 GEMM + softmax 组合 attention", "Phase4 暂不做 Flash Attention，先做标准组合算子"],
        ],
    )
    h(doc, "6.1 require_input_count", 2)
    p(doc, "这个 helper 不做计算，只做防御式检查。比如 MatMul 至少需要两个输入，Attention 至少需要三个输入。如果 JSON 或手写图错了，它能给出明确 error 字符串。")
    h(doc, "6.2 execute_matmul：Phase4 的主要加速点", 2)
    p(doc, "Phase3 的 MatMul 调用 gemm_naive_ikj，Phase4 改成 gemm_parallel_tiled。这个函数来自 Phase1，结合 cache blocking 和 OpenMP 多线程。")
    code(doc, """Phase3:
gemm_naive_ikj(M, N, K, A, B, C)

Phase4:
gemm_parallel_tiled(M, N, K, A, B, output.data)""")
    p(doc, "这个改动体现 Phase4 的核心目标：图结构不变，底层算子替换为更快实现。")
    h(doc, "6.3 execute_softmax：为什么按行并行", 2)
    p(doc, "softmax_2d 的每一行独立。第 r 行只读写 matrix[r*cols : (r+1)*cols)。所以可以用 OpenMP parallel for 把不同 row 分配给不同线程，不会发生写冲突。")
    h(doc, "6.4 execute_layernorm：Phase4 修正为 row-wise", 2)
    p(doc, "Transformer 中输入通常是 [seq_len, d_model]。LayerNorm 应该对每个 token 的 d_model 维度单独归一化，而不是对整个 seq_len*d_model 展平归一化。Phase4 用 cols=in.shape.back()，rows=data.size()/cols 来实现这个语义。")
    code(doc, """for each row r:
    layernorm(in[r, :], cols, eps, gamma, beta, out[r, :])""")
    h(doc, "6.5 execute_gelu：为什么用 approx", 2)
    p(doc, "Phase3 用 gelu_vector_exact，优先正确性；Phase4 用 gelu_approx，符合推理场景常见取舍：用 tanh 近似减少 erf 的计算成本。")
    h(doc, "6.6 execute_attention：标准 Attention 组合", 2)
    p(doc, "Phase4 的 Attention 仍然是标准三步，不是 Flash Attention。它先并行转置 K，再用 parallel GEMM 算 scores，再对 scores 乘 scale，再逐行 softmax，最后 parallel GEMM 乘 V。")
    code(doc, """K_T[j * seq_len + i] = K[i * d_k + j]
scores = Q x K_T
scores *= 1 / sqrt(d_k)
softmax each row of scores
output = scores x V""")
    p(doc, "这条路径很清晰，也方便 benchmark。Phase5 可以在这里替换成 fused/Flash Attention。")


def add_phase4_parallel_executor(doc: Document) -> None:
    h(doc, "7. Phase4：parallel_graph_executor.h / .cpp", 1)
    p(doc, "ParallelGraphExecutor 负责把 Phase3 的线性拓扑序进一步分成 level。同一 level 内的节点互相没有依赖，因此可以并行执行。")
    table(
        doc,
        ["函数/结构", "作用", "关键点"],
        [
            ["JobResult", "std::async 返回值", "携带 node、output、ok、error，主线程统一写回 graph"],
            ["gather_inputs", "从 graph 收集某个 node 的输入 Tensor", "当前实现复制 Tensor，避免多线程读写 unordered_map 的风险"],
            ["ParallelGraphExecutor::ParallelGraphExecutor", "保存 backend 引用和 verbose 开关", "同一个执行器可跑 serial 或 level-parallel"],
            ["execute_serial", "按 sorted_nodes 顺序执行 backend", "Phase4 版线性调度，用于和 level-parallel 对照"],
            ["build_execution_levels", "把拓扑序转成 level 列表", "计算每个节点所在层级"],
            ["execute_level_parallel", "每层内用 std::async 并发执行", "每层结束后统一 set_tensor"],
        ],
    )
    h(doc, "7.1 为什么要 gather_inputs 复制 Tensor", 2)
    p(doc, "ComputeGraph 内部用 unordered_map 保存 Tensor。unordered_map 不是线程安全容器。如果多个线程一边读一边写 graph，可能出错。所以当前实现先在主线程把本 level 所有输入复制出来，worker 线程只处理自己的副本，最后主线程统一写回输出。")
    p(doc, "这会带来额外拷贝，所以 benchmark 中 level-parallel 可能比 serial 慢。这是正确性优先的版本。后续可以引入只读 TensorView、线程安全 TensorStore 或引用计数 buffer 来减少拷贝。")
    h(doc, "7.2 build_execution_levels 的逻辑", 2)
    p(doc, "它维护 producer_level：某个 Tensor 名是在哪一层被生产的。节点的 level 等于它所有输入生产层级的最大值 + 1。如果输入不是图内生产的，比如 x、Wq，它就是外部输入，不影响 level。")
    code(doc, """level(node) = max(level(producer(input)) + 1 for graph-produced inputs)
if no graph-produced input:
    level(node) = 0""")
    p(doc, "对于 attention_block，q_proj、k_proj、v_proj 都只依赖外部输入 x 和权重，因此都在 level 0；attn_out 依赖它们，所以在 level 1；ln_out 依赖 attn_out，所以在 level 2。")
    h(doc, "7.3 execute_level_parallel 的执行顺序", 2)
    num(doc, "先调用 build_execution_levels。")
    num(doc, "对每一层，先 gather_inputs，主线程完成输入快照。")
    num(doc, "对该层每个 node 启动一个 std::async 任务。")
    num(doc, "每个任务调用 backend_.execute。")
    num(doc, "主线程 future.get 收集 JobResult。")
    num(doc, "如果失败，打印 node id 和 error，返回 false。")
    num(doc, "如果成功，先保存 output_name，再 std::move(output) 写回 graph。")
    p(doc, "保存 output_name 这一步非常重要：C++ 函数实参求值顺序可能导致 std::move(output) 先发生，如果直接写 graph.set_tensor(output.name, std::move(output))，output.name 可能已经被移动，导致 Tensor 写入空名字。这是 Phase4 实现中特别值得记住的 C++ 细节。")


def add_phase4_quant(doc: Document) -> None:
    h(doc, "8. Phase4：quantization.h / quantization.cpp", 1)
    p(doc, "量化模块实现的是 weight-only INT8 GEMM：激活 A 仍是 FP32，权重 B 量化为 INT8，计算时把 INT8 乘回 scale 得到近似 FP32 结果。")
    table(
        doc,
        ["函数/结构", "作用", "细节"],
        [
            ["QuantizedTensor", "保存量化后的张量", "shape、int8 values、scale 列表、per_channel 标志"],
            ["clamp_to_i8", "把整数截断到 [-127,127]", "对称量化通常避免使用 -128"],
            ["safe_scale", "避免 max_abs=0 时除零", "全零张量 scale=1"],
            ["quantize_symmetric_per_tensor", "整个 Tensor 共用一个 scale", "实现简单，误差可能较大"],
            ["quantize_weight_per_output_channel", "每个输出通道/列一个 scale", "权重量化常用方式，精度更好"],
            ["gemm_fp32_int8_weight", "A_fp32 x B_int8 -> C_fp32", "累加时乘 int8，最后乘 scale"],
            ["max_abs_error", "计算两个结果最大绝对误差", "测试和 benchmark 用它衡量量化误差"],
        ],
    )
    h(doc, "8.1 对称量化公式", 2)
    code(doc, """scale = max(abs(W)) / 127
W_int8 = round(W_fp32 / scale)
W_fp32_approx = W_int8 * scale""")
    p(doc, "per-tensor 量化只有一个 scale，所有元素共享。per-channel 量化每一列一个 scale，更适合 GEMM 权重，因为每个输出通道的数值范围可能不同。")
    h(doc, "8.2 gemm_fp32_int8_weight 的计算", 2)
    p(doc, "函数输入 A 是 [M,K] 的 FP32，B 是 [K,N] 的 INT8 权重，输出 C 是 [M,N] 的 FP32。")
    code(doc, """for i in M:
  for j in N:
    acc = 0
    for k in K:
      acc += A[i,k] * B_int8[k,j]
    C[i,j] = acc * scale[j]""")
    p(doc, "当前实现是标量循环加 OpenMP，并没有使用 AVX2/VNNI 指令，所以它的主要价值是建立量化正确性和误差评估链路。后续可以把内层循环替换成 SIMD INT8 kernel。")


def add_tests_bench_cmake(doc: Document) -> None:
    h(doc, "9. 测试、Benchmark 和 CMake", 1)
    h(doc, "9.1 Phase4 测试文件 test_phase4_cpu.cpp", 2)
    table(
        doc,
        ["测试函数", "验证什么", "为什么重要"],
        [
            ["test_execution_levels", "attention graph 被分成 3 层，第一层有 q/k/v 三个节点", "证明 levelization 符合 DAG 依赖"],
            ["test_serial_matches_level_parallel", "serial 和 level-parallel 输出一致", "证明并行调度没有改变计算语义"],
            ["test_rowwise_layernorm", "每一行 LayerNorm 后均值约 0、方差约 1", "证明 Phase4 LayerNorm 是 transformer 语义"],
            ["test_quantized_weight_gemm", "INT8 weight GEMM 与 FP32 GEMM 的 max error < 0.05", "证明量化路径误差可控"],
        ],
    )
    h(doc, "9.2 Benchmark 文件 bench_phase4_cpu.cpp", 2)
    p(doc, "benchmark 分两组：GEMM Paths 比较 naive、parallel、int8w；Graph Schedule 比较 Phase4 serial schedule 和 level-parallel schedule。time_ms 运行多次取最小值，减少偶发抖动影响。")
    p(doc, "当前实测中 512 方阵 parallel GEMM 明显快于 naive；INT8 weight-only 在小尺寸有收益，但 512 上慢于 parallel FP32，因为当前没有使用专门的 AVX2/VNNI INT8 指令。level-parallel 慢于 serial，是因为输入快照复制和 OpenMP 线程叠加造成额外开销。")
    h(doc, "9.3 CMake 文件怎么把模块组合起来", 2)
    p(doc, "Phase3 CMake 构建 graph_engine，并链接 Phase1 GEMM 与 Phase2 operators。Phase4 CMake 构建 phase4_cpu_accel，并复用 Phase1/2/3 代码。")
    code(doc, """phase4_cpu_accel links:
  graph_core      -> Phase3 graph/json_parser/memory_pool
  operators       -> Phase2 softmax/layernorm/gelu/attention
  gemm_phase1     -> Phase1 naive/tiled/parallel GEMM
  OpenMP::OpenMP_CXX
  Threads::Threads""")


def add_combined_flow(doc: Document) -> None:
    h(doc, "10. Phase3 与 Phase4 如何组合运行", 1)
    p(doc, "Phase4 没有替换 ComputeGraph。它复用 Phase3 的图结构，只替换执行策略和底层算子。")
    table(
        doc,
        ["步骤", "Phase3 做什么", "Phase4 做什么"],
        [
            ["构图", "JsonParser / add_node / set_tensor", "直接复用"],
            ["排序", "ComputeGraph::topological_sort", "直接复用 sorted_nodes"],
            ["执行策略", "Executor::execute 线性 if-else", "ParallelGraphExecutor 可 serial 或 level-parallel"],
            ["算子实现", "Phase1 naive GEMM + Phase2 exact operators", "OpenMP parallel GEMM/Softmax/LayerNorm/GELU/Attention"],
            ["内存", "MemoryPool 教学型复用", "当前先用 Tensor copy 保证线程安全，后续可升级 TensorView/arena"],
            ["量化", "无", "新增 INT8 weight-only GEMM"],
        ],
    )
    h(doc, "10.1 对 attention_block 的逐节点执行", 2)
    code(doc, """输入:
  x  = [seq_len, d_model]
  Wq = [d_model, d_model]
  Wk = [d_model, d_model]
  Wv = [d_model, d_model]

level 0:
  q_proj = MatMul(x, Wq)
  k_proj = MatMul(x, Wk)
  v_proj = MatMul(x, Wv)

level 1:
  attn_out = Attention(q_proj, k_proj, v_proj)

level 2:
  ln_out = LayerNorm(attn_out)""")
    p(doc, "Phase3 会把它排成合法线性序列；Phase4 会进一步发现 q/k/v 三个投影互不依赖，因此可以在同一层并发。")
    h(doc, "10.2 你应该如何调试这个框架", 2)
    num(doc, "先 dump graph，确认节点 id、op_type、inputs、outputs 是否正确。")
    num(doc, "检查 topological_sort 的 sorted_nodes 数量是否等于 node_count。")
    num(doc, "如果执行失败，看 error 是缺 Tensor、维度不匹配，还是 unknown op。")
    num(doc, "对 MatMul，永远先写出 A=[M,K]、B=[K,N]、C=[M,N]。")
    num(doc, "对 Attention，确认 Q/K/V 的 seq_len 一致，Q/K 的 d_k 一致。")
    num(doc, "对并行执行，先和 serial 输出做 max_diff 比较，再看性能。")
    num(doc, "对量化，先看 max_abs_error，再谈速度。")


def add_learning_path(doc: Document) -> None:
    h(doc, "11. 按代码阅读的推荐顺序", 1)
    num(doc, "读 graph.h：理解 Tensor、Node、ComputeGraph 三个核心概念。")
    num(doc, "读 graph.cpp：理解 add_node、set_tensor、get_tensor、topological_sort。")
    num(doc, "读 json_parser.cpp：理解 JSON 如何变成 Node。")
    num(doc, "读 executor.cpp 的 execute_matmul：掌握一个节点执行的标准模板。")
    num(doc, "读 executor.cpp 的 execute：理解 Phase3 线性调度。")
    num(doc, "读 memory_pool.cpp：理解 allocate/deallocate/compaction。")
    num(doc, "读 backend.h：理解 Phase4 为什么需要后端抽象。")
    num(doc, "读 cpu_parallel_backend.cpp：逐个对照 Phase3 算子，看它们怎样变成 OpenMP 版本。")
    num(doc, "读 parallel_graph_executor.cpp：理解 level 并行、std::async、JobResult、线程安全写回。")
    num(doc, "读 quantization.cpp：理解 scale、int8 values、per-channel GEMM。")
    num(doc, "最后读 tests 和 benchmark：学习怎样证明正确性和怎样解释性能数据。")
    h(doc, "11.1 最重要的工程认知", 2)
    bullet(doc, "计算图是骨架，算子是血肉。Phase4 的后端替换不应该破坏 Phase3 的图表示。")
    bullet(doc, "拓扑排序解决“能不能按依赖顺序执行”，levelization 解决“哪些节点可以同时执行”。")
    bullet(doc, "并行不是一定更快。复制开销、线程启动、OpenMP 过度订阅都会吞掉收益。")
    bullet(doc, "量化不是只看速度，还必须给出误差。没有误差报告的 INT8 benchmark 没有意义。")
    bullet(doc, "Phase4 当前是正确、可测、可讲清楚的 CPU 加速 MVP；真正工业级优化可以继续做 SIMD INT8、TensorView、内存池零拷贝和 CPU Flash Attention。")


def main() -> None:
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_overview(doc)
    add_phase3_graph(doc)
    add_phase3_json(doc)
    add_phase3_memory(doc)
    add_phase3_executor(doc)
    add_phase4_backend(doc)
    add_phase4_parallel_executor(doc)
    add_phase4_quant(doc)
    add_tests_bench_cmake(doc)
    add_combined_flow(doc)
    add_learning_path(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
